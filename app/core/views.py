# views.py
import os
import subprocess
from datetime import datetime, timedelta
from docx import Document
from django.http import FileResponse, JsonResponse
from rest_framework.views import APIView
from rest_framework.parsers import JSONParser

class CoreView(APIView):
    parser_classes = [JSONParser]
    
    def post(self, request):
        dados = request.data
        
        # Processamento de valores e datas
        try:
            valor_atualizacao = float(dados["valor-atualizacao"])
            nro_parcelas = int(dados["nro-parcelas"])
            conta_parcelas = valor_atualizacao / nro_parcelas
            dados["conta-parcelas"] = f"{conta_parcelas:.2f}"
        except ValueError:
            return JsonResponse({"error": "Erro no cálculo de parcelas."}, status=400)

        try:
            data_primeiro_pgto = datetime.strptime(dados["data-primeiro-pagamento"], "%d/%m/%Y")
            data_primeiro_pgto_30 = data_primeiro_pgto + timedelta(days=30)
            dados["data-primeiro-pagamento-30"] = data_primeiro_pgto_30.strftime("%d/%m/%Y")
        except ValueError:
            return JsonResponse({"error": "Erro no processamento da data."}, status=400)

        # Definição do modelo de contrato
        arquivo_modelo_sim = "./modelosim.docx"
        arquivo_modelo_nao = "./modelonao.docx"
        arquivo_entrada = arquivo_modelo_sim if dados.get("tem-manutencao-mensal", "nao").lower() == "sim" else arquivo_modelo_nao
        
        arquivo_saida_docx = "./contrato_temp.docx"
        arquivo_saida_pdf = "./contrato_final.pdf"

        # Preenchimento do modelo
        try:
            doc = Document(arquivo_entrada)
            for paragrafo in doc.paragraphs:
                for run in paragrafo.runs:
                    for chave, valor in dados.items():
                        if f"<{chave}>" in run.text:
                            run.text = run.text.replace(f"<chave>", valor)

            for tabela in doc.tables:
                for linha in tabela.rows:
                    for celula in linha.cells:
                        for paragrafo in celula.paragraphs:
                            for run in paragrafo.runs:
                                for chave, valor in dados.items():
                                    if f"<{chave}>" in run.text:
                                        run.text = run.text.replace(f"<chave>", valor)
            
            doc.save(arquivo_saida_docx)
        except Exception as e:
            return JsonResponse({"error": f"Erro ao processar documento: {str(e)}"}, status=500)
        
        # Conversão para PDF
        try:
            subprocess.run(["soffice", "--headless", "--convert-to", "pdf", arquivo_saida_docx, "--outdir", "./"], check=True)
            os.remove(arquivo_saida_docx)
        except Exception as e:
            return JsonResponse({"error": f"Erro ao converter para PDF: {str(e)}"}, status=500)
        
        return FileResponse(open(arquivo_saida_pdf, "rb"), as_attachment=True, filename="contrato.pdf")
